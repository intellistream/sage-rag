"""Microsoft Word document loaders (.docx / .doc).

Moved from sage.libs.rag.document_loaders — concrete implementation belongs in sage-rag.

- DocxLoader  requires: python-docx  (cross-platform)
- DocLoader   requires: pywin32      (Windows only; use DocxLoader on Linux/Mac)
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from sage_rag.interface import Document, DocumentLoader


class DocxLoader(DocumentLoader):
    """Load modern Word documents (.docx).

    Note:
        Requires ``python-docx``::

            pip install python-docx

    Example:
        >>> from sage_libs.sage_rag import DocxLoader
        >>> loader = DocxLoader()
        >>> doc = loader.load("report.docx")
    """

    def load(self, source: str, **kwargs: Any) -> Document:
        """Load a .docx file.

        Args:
            source: Path to the Word document.
            **kwargs: Additional options.

        Returns:
            Document with extracted text.

        Raises:
            FileNotFoundError: If file does not exist.
            ImportError: If python-docx is not installed.
        """
        try:
            import docx  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError(
                "DocxLoader requires 'python-docx'. Install with: pip install python-docx"
            ) from exc

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")

        document = docx.Document(str(path))
        content = "\n".join(para.text for para in document.paragraphs)

        return Document(
            content=content,
            metadata={
                "source": str(path.absolute()),
                "filename": path.name,
                "extension": path.suffix,
                "size_bytes": path.stat().st_size,
                "loader": "docx",
            },
        )

    def load_batch(self, sources: list[str], **kwargs: Any) -> list[Document]:
        """Load multiple .docx files."""
        return [self.load(source, **kwargs) for source in sources]

    def supported_formats(self) -> list[str]:
        return [".docx"]


class DocLoader(DocumentLoader):
    """Load legacy Word documents (.doc) via the Windows COM interface.

    .. warning::
        This loader is **Windows-only**. On Linux/macOS, convert the file to
        ``.docx`` first (e.g. with LibreOffice) and use :class:`DocxLoader`.

    Note:
        Requires ``pywin32``::

            pip install pywin32  # Windows only

    Example:
        >>> from sage_libs.sage_rag import DocLoader
        >>> loader = DocLoader()
        >>> doc = loader.load("legacy.doc")
    """

    def load(self, source: str, **kwargs: Any) -> Document:
        """Load a .doc file via the Windows COM interface.

        Args:
            source: Path to the legacy Word document.
            **kwargs: Additional options.

        Returns:
            Document with extracted text.

        Raises:
            FileNotFoundError: If file does not exist.
            ImportError: If pywin32 is not installed or not on Windows.
        """
        try:
            import win32com.client  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError(
                "DocLoader requires 'pywin32' (Windows only). "
                "Install with: pip install pywin32  "
                "On Linux/macOS, convert to .docx first."
            ) from exc

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(path.resolve()))
            content: str = doc.Content.Text
        finally:
            doc.Close()  # type: ignore[possibly-undefined]
            word.Quit()

        return Document(
            content=content,
            metadata={
                "source": str(path.absolute()),
                "filename": path.name,
                "extension": path.suffix,
                "size_bytes": path.stat().st_size,
                "loader": "doc",
            },
        )

    def load_batch(self, sources: list[str], **kwargs: Any) -> list[Document]:
        """Load multiple .doc files."""
        return [self.load(source, **kwargs) for source in sources]

    def supported_formats(self) -> list[str]:
        return [".doc"]
